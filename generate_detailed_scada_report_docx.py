import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color_hex):
    """Applies background color shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) of a cell in DXA (1/20th of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color_hex="CBD5E1"):
    """Applies clean slate-colored borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_styled_heading(doc, text, level, space_before=12, space_after=4):
    """Adds a styled heading with standard microgrid colors."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Blue
        p.paragraph_format.space_before = Pt(16)
    else:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate Dark
        p.paragraph_format.space_before = Pt(10)
        
    return p

def format_cell_text(cell, text, bold=False, color_rgb=(0x33, 0x41, 0x55), size=8.0, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Sets standard formatting for paragraphs inside table cells."""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

def add_flowchart_table(doc):
    """Inserts a styled visual flow diagram representing the 1Hz microgrid controller execution loop."""
    flow_steps = [
        ("Step 1: Start 1Hz Loop Tick", [
            "• Background daemon loop triggers asynchronously every 1.0 second.",
            "• Aligns timer offsets and initializes temporary local registers."
        ]),
        ("Step 2: Telemetry Data Scanning & Ingestion", [
            "• Reads Dropdown connection status configurations for each microgrid asset.",
            "• Fallback (Sample CSV): Loosely reads next sequential row cache from CSV files.",
            "• Connected (Live Protocol): Scans simulated registers (Modbus TCP, CAN Bus, OPC UA, IEC)."
        ]),
        ("Step 3: Diagnostics & Asset RUL Estimations", [
            "• Runs Isolation Forest anomaly scores over grid frequency, current, voltage, and efficiency.",
            "• Evaluates heatsink thermal stresses and battery capacity fade parameters.",
            "• Projects remaining useful life (RUL) limits in hours for components."
        ]),
        ("Step 4: EMS Dispatch Controller Decision Matrix", [
            "• Checks safety bounds: Battery SOC limits (20% - 95%) and temperature derating.",
            "• Connected Mode: Renewables priority -> BESS charge -> Grid export.",
            "• Outage Mode: Battery discharge -> Tiered load-shedding levels (30% / 60% / 85%)."
        ]),
        ("Step 5: Logging & Alarm DB Commits", [
            "• Writes 1Hz time-series snapshots to SQLite scada_historian table.",
            "• Evaluates safety bounds: registers warning/critical status warnings in alarms_log."
        ]),
        ("Step 6: Fieldbus Map updates & WebSockets Broadcast", [
            "• Encodes CAN Bus BMS packages and updates Modbus registers starting at 40001.",
            "• Broadcasts real-time telemetry JSON payload to browser WebSocket clients."
        ])
    ]
    
    for i, (title, details) in enumerate(flow_steps):
        # Create a single-cell box for the step
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        cell.width = Inches(5.5)
        set_cell_shading(cell, "F0F9FF") # light blue shade (blue-50)
        set_cell_margins(cell, top=70, bottom=70, left=120, right=120)
        
        # Apply borders to the box
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for b_name in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{b_name}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '6')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '3B82F6') # Blue border
            tcBorders.append(b)
        tcPr.append(tcBorders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run_title = p.add_run(title)
        run_title.font.name = 'Arial'
        run_title.font.bold = True
        run_title.font.size = Pt(8.5)
        run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        for d in details:
            p_det = cell.add_paragraph()
            p_det.paragraph_format.space_before = Pt(0)
            p_det.paragraph_format.space_after = Pt(1)
            p_det.paragraph_format.line_spacing = 1.05
            run_det = p_det.add_run(d)
            run_det.font.name = 'Calibri'
            run_det.font.size = Pt(7.5)
            run_det.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            
        # Draw arrow to the next box (except the last one)
        if i < len(flow_steps) - 1:
            p_arr = doc.add_paragraph()
            p_arr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arr.paragraph_format.space_before = Pt(2)
            p_arr.paragraph_format.space_after = Pt(2)
            run_arr = p_arr.add_run("↓")
            run_arr.font.name = 'Calibri'
            run_arr.font.size = Pt(12)
            run_arr.font.bold = True
            run_arr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8) # light gray arrow

def build_docx_report():
    doc = Document()
    
    # 1. Page Margins Setup (0.5 in margins to match PDF)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # 2. Document Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("APEX MICROGRID SCADA & EMS SYSTEM REPORT")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Blue
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("Comprehensive Technical Specifications, Dataset Analytics & Working Architecture")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(10)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6) # Light Blue
    
    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, "E2E8F0")
    
    metadata = [
        ("Project Name", "APEX EMS — Hybrid Microgrid Energy Management System & SCADA Platform"),
        ("Project Type", "Supervisory Control, High-Frequency Logger, AI Optimizer & Fault Advisor"),
        ("Execution Rate", "1Hz Non-Blocking Continuous Daemon Loop (1-second intervals)"),
        ("Communications", "WebSockets Broadcasts, REST APIs, Modbus TCP, CAN Bus, OPC UA, IEC 61850")
    ]
    
    col_widths_meta = [Inches(1.5), Inches(6.0)]
    for i, (label, val) in enumerate(metadata):
        row = meta_table.rows[i]
        row.cells[0].width = col_widths_meta[0]
        row.cells[1].width = col_widths_meta[1]
        
        set_cell_margins(row.cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(row.cells[1], top=80, bottom=80, left=100, right=100)
        
        format_cell_text(row.cells[0], label, bold=True, color_rgb=(0x1E, 0x3A, 0x8A), size=8.5)
        format_cell_text(row.cells[1], val, size=8.5)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # 1. Project Overview & Scope
    add_styled_heading(doc, "1. Project Overview & Scope", level=1)
    p_overview = doc.add_paragraph()
    p_overview.paragraph_format.space_after = Pt(6)
    p_overview.paragraph_format.line_spacing = 1.15
    run_overview = p_overview.add_run(
        "The APEX Microgrid SCADA & EMS is a production-grade supervisory control and data acquisition system integrated "
        "with a rule-based state machine and reinforcement learning AI agent. The platform coordinates a hybrid electrical "
        "network consisting of Solar PV generation, Wind turbine generation, a Battery Energy Storage System (BESS), utility "
        "grid coupling, and dynamic facility loads. Operating at a 1Hz execution frequency, the background daemon handles "
        "telemetry ingestion, executes safety checks, runs forecasting models, updates simulated physics, and dispatches "
        "optimized charge/discharge profiles. It ensures microgrid stability, minimizes tariff-based import fees, protects "
        "battery chemistry under temperature excursions, and generates structured alarms for immediate field restoration."
    )
    run_overview.font.name = 'Calibri'
    run_overview.font.size = Pt(8.5)
    run_overview.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # 2. Core Requirements
    add_styled_heading(doc, "2. Core Requirements", level=1)
    
    reqs = [
        ("High-Fidelity Synoptic Visualization", "Render browser-based HMI dashboards mapping active power flows with dynamic speed animation, colored tag displays, and diagnostic faceplates."),
        ("Continuous Replay Loop Fallback", "Loop default sample CSV datasets sequentially at 1Hz if sensors are disconnected or configured to 'Sample CSV'. Connection to protocols promotes metrics automatically to live simulation."),
        ("AI In-situ Optimization", "Utilize reinforcement learning (Q-Learning) and machine learning forecasters (Neural Networks, Gradient Boosting) pre-trained on datasets to schedule battery dispatch cost-effectively."),
        ("Industrial Protocol Mapping", "Interface physical asset tags with standardized automation layers including Modbus TCP registers, CAN Bus BMS binary frames, OPC UA structures, and IEC 61850 Logical Nodes."),
        ("Fault Diagnostics & Advisory System", "Deploy Isolation Forest anomaly detection to isolate sensor drifts, evaluate remaining useful life (RUL) of components, and pair warnings with structured repair actions.")
    ]
    
    for title, desc in reqs:
        p_req = doc.add_paragraph(style='List Bullet')
        p_req.paragraph_format.space_after = Pt(3)
        p_req.paragraph_format.left_indent = Inches(0.25)
        
        run_t = p_req.add_run(f"{title}: ")
        run_t.font.name = 'Calibri'
        run_t.font.size = Pt(8.5)
        run_t.font.bold = True
        run_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        run_d = p_req.add_run(desc)
        run_d.font.name = 'Calibri'
        run_d.font.size = Pt(8.5)
        run_d.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # 3. Technology Stack Table
    add_styled_heading(doc, "3. Technology Stack", level=1)
    
    tech_headers = ["Layer", "Technology", "Purpose & Implementation"]
    tech_rows = [
        ["Frontend UI", "React 19, TailwindCSS, Chart.js, Vite", "Renders responsive HMI dashboards, connection widgets, logs, and alarms journal views."],
        ["Vanilla HMI", "Vanilla Javascript, HTML5 Canvas, HSL CSS", "Legacy browser dashboard displaying dynamic synoptic SVG flow particles and popups."],
        ["Backend Core", "FastAPI (Python 3.10+), Uvicorn", "Asynchronous web server hosting REST controller overrides routing and WebSockets telemetry."],
        ["Historian DB", "SQLite (PostgreSQL compatible)", "Maintains configurations settings, alarm logs journal, and high-frequency time-series tables."],
        ["AI & Diagnostic", "Scikit-Learn, MLP Neural Networks, Isolation Forest", "Predicts load, solar capacities, detects anomalous metrics, and calculates RUL indexes."],
        ["Automation", "Modbus TCP, CAN Bus, OPC UA, IEC 61850", "Simulates holding registers, 8-byte BMS CAN frames, OPC UA structures, and circuit breakers."]
    ]
    
    tech_table = doc.add_table(rows=len(tech_rows) + 1, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tech_table)
    
    hdr_row = tech_table.rows[0]
    col_widths_tech = [Inches(1.2), Inches(2.0), Inches(4.3)]
    for j, text in enumerate(tech_headers):
        cell = hdr_row.cells[j]
        cell.width = col_widths_tech[j]
        set_cell_shading(cell, "1E3A8A")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        format_cell_text(cell, text, bold=True, color_rgb=(0xFF, 0xFF, 0xFF), size=8.0)
        
    for i, row_data in enumerate(tech_rows):
        row = tech_table.rows[i + 1]
        bg_color = "FFFFFF" if i % 2 == 0 else "F8FAFC"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths_tech[j]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            is_bold = (j == 0)
            color = (0x1E, 0x3A, 0x8A) if is_bold else (0x33, 0x41, 0x55)
            format_cell_text(cell, val, bold=is_bold, color_rgb=color, size=7.5)
            
    # 4. System Architecture
    add_styled_heading(doc, "4. System Architecture", level=1)
    p_arch = doc.add_paragraph()
    p_arch.paragraph_format.space_after = Pt(6)
    p_arch.paragraph_format.line_spacing = 1.15
    run_arch = p_arch.add_run(
        "The APEX platform follows a decoupled client-server architecture. The browser client connects to the FastAPI "
        "gateway using REST for settings changes and persistent WebSockets for live telemetry data streams. The backend "
        "runs an independent background loop mapping simulated microgrid states to databases and industrial protocol "
        "drivers. The following diagram illustrates the flow and organization of core microgrid services and assets:"
    )
    run_arch.font.name = 'Calibri'
    run_arch.font.size = Pt(8.5)
    run_arch.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # Insert System Architecture Image
    try:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(12)
        r_img = p_img.add_run()
        r_img.add_picture("system_architecture.png", width=Inches(6.5))
    except Exception as e:
        print(f"[DOCX Compile Warning] Failed to insert system_architecture.png: {e}")
        
    doc.add_page_break()
    
    # ================= PAGE 2 =================
    add_styled_heading(doc, "5. Telemetry Dataset Architecture", level=1)
    
    p_ds_intro = doc.add_paragraph()
    p_ds_intro.paragraph_format.space_after = Pt(6)
    run_ds_intro = p_ds_intro.add_run(
        "The APEX system preloads default historical datasets to pretrain machine learning models and to drive "
        "dashboard telemetry tags during fallback mode when hardware sensors are disconnected or configured to 'Sample CSV'."
    )
    run_ds_intro.font.name = 'Calibri'
    run_ds_intro.font.size = Pt(8.5)
    run_ds_intro.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    dataset_headers = ["Dataset File", "Primary Columns", "Physical Dynamics & Formulas", "Role & Size"]
    dataset_rows = [
        ["solar_dataset.csv", "timestamp, Solar_Power, Solar_Voltage, Solar_Current, Solar_Temperature", 
         "100 kW max capacity. Sinusoidal power peak at 12:00 PM: P = 100 * sin(pi * (h-6)/12). Panels heat up under sunlight: T = T_ambient + 15 * IrradFactor. Volt = 380-420V DC.", "202 rows. Pretrains solar models and fallback loops."],
        
        ["wind_dataset.csv", "timestamp, Wind_Power, Wind_Speed, Wind_RPM", 
         "50 kW max capacity. Cut-in: 3.0 m/s, Rated: 12.0 m/s, Cut-out: 25.0 m/s. Cubic curve P = 50 * ((v-3)/9)^3. RPM scales: RPM = 100 + 400 * ((v-3)/9).", "202 rows. Pretrains wind turbine models."],
        
        ["battery_dataset.csv", "timestamp, Battery_SOC, Battery_Voltage, Battery_Current, Battery_Temp, Battery_SOH", 
         "200 kWh BESS pack. 96S configuration. Pack voltage: V = 320.0 + 95.0 * (SOC/100) + I * 0.15. Charge current is negative; discharge is positive.", "202 rows. Pretrains BESS SOC prediction model."],
        
        ["load_dataset.csv", "timestamp, Load_Demand, Load_Voltage, Load_Current", 
         "Dynamic load with dual daily peaks (8 AM and 7 PM). Base load of 25 kW. Single phase AC feeder: V = 230V, I = (P * 1000) / V.", "202 rows. Pretrains auto-regressive MLP load forecaster."],
        
        ["inverter_dataset.csv", "timestamp, Inverter_Output_Power, Inverter_Efficiency, Inverter_Status", 
         "Conversion efficiency scales from 95% to 99%, dipping at low capacities due to core losses: Eff = 98.5 - (1.5 / LoadRatio).", "202 rows. Drives HMI in default telemetry mode loops."],
        
        ["grid_dataset.csv", "timestamp, Grid_Power, Grid_Voltage, Grid_Frequency", 
         "Nominal 400V 3-phase connection at 50 Hz. Positive = Import from utility, Negative = Export. Frequency drifts slightly with network load.", "202 rows. Grid connection time-series fallback logs."],
        
        ["grid_exports_dataset.csv", "datetime, timestamp, solar_power, wind_power, battery_soc, grid_power, load_demand, ems_action, electricity_cost...", 
         "Log-traces all 26 telemetry points in the microgrid system at 1Hz execution frequency. Used to validate reinforcement learning optimizer.", "13,422 rows (2.67 MB). Audits dispatch, cost analytics & neural nets."]
    ]
    
    ds_table = doc.add_table(rows=len(dataset_rows) + 1, cols=4)
    ds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(ds_table)
    
    col_widths_ds = [Inches(1.2), Inches(1.3), Inches(3.7), Inches(1.3)]
    
    hdr_row = ds_table.rows[0]
    for j, text in enumerate(dataset_headers):
        cell = hdr_row.cells[j]
        cell.width = col_widths_ds[j]
        set_cell_shading(cell, "1E3A8A")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        format_cell_text(cell, text, bold=True, color_rgb=(0xFF, 0xFF, 0xFF), size=8.0)
        
    for i, row_data in enumerate(dataset_rows):
        row = ds_table.rows[i + 1]
        bg_color = "FFFFFF" if i % 2 == 0 else "F8FAFC"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths_ds[j]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            is_bold = (j == 0)
            color = (0x1E, 0x3A, 0x8A) if is_bold else (0x33, 0x41, 0x55)
            format_cell_text(cell, val, bold=is_bold, color_rgb=color, size=7.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # 6. Key Parameters & Threshold Limits
    add_styled_heading(doc, "6. Key Parameters & Threshold Limits", level=1)
    p_lim_intro = doc.add_paragraph()
    p_lim_intro.paragraph_format.space_after = Pt(6)
    run_lim_intro = p_lim_intro.add_run(
        "SCADA supervisors continuously evaluate measurements against physical limits to trigger safety overrides "
        "and register alarms in the Event Journal."
    )
    run_lim_intro.font.name = 'Calibri'
    run_lim_intro.font.size = Pt(8.5)
    run_lim_intro.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    thresh_headers = ["Parameter", "Normal Range", "Warning Limit", "Critical Limit", "Mitigation Action"]
    thresh_rows = [
        ["Battery SOC", "20.0% - 95.0%", "SOC <= 20.0%", "SOC <= 15.0%", "Activates tiered load shedding in islanding mode to protect battery cells."],
        ["Battery Temp", "15.0C - 35.0C", "Temp >= 45.0C", "Temp >= 55.0C / Temp <= -10.0C", "Warning derates charge rates by 50%. Critical isolates BESS (0 kW flow)."],
        ["Feeder Current", "10A - 150A", "Current >= 160A", "Current >= 180A", "Triggers over-current warning (LOAD-OC-001) and opens breakers."],
        ["Grid Voltage", "380V - 420V", "Voltage >= 430V", "Voltage >= 440V", "Registers grid overvoltage fault (GRID-OV-001). Isolates inverter."],
        ["Grid Frequency", "49.8 - 50.2 Hz", "Freq <= 49.5 / >= 50.5 Hz", "Freq <= 49.0 / >= 51.0 Hz", "Registers warning alarm. Triggers utility sync mismatch warning."],
        ["Inverter Eff", "90.0% - 99.0%", "Eff <= 88.0%", "Eff <= 85.0%", "Flags inverter thermal cooling or IGBT wear anomalies in diagnostic suite."]
    ]
    
    thresh_table = doc.add_table(rows=len(thresh_rows) + 1, cols=5)
    thresh_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(thresh_table)
    
    col_widths_thresh = [Inches(1.1), Inches(1.1), Inches(1.3), Inches(1.3), Inches(2.7)]
    
    hdr_row = thresh_table.rows[0]
    for j, text in enumerate(thresh_headers):
        cell = hdr_row.cells[j]
        cell.width = col_widths_thresh[j]
        set_cell_shading(cell, "1E3A8A")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        format_cell_text(cell, text, bold=True, color_rgb=(0xFF, 0xFF, 0xFF), size=8.0)
        
    for i, row_data in enumerate(thresh_rows):
        row = thresh_table.rows[i + 1]
        bg_color = "FFFFFF" if i % 2 == 0 else "F8FAFC"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths_thresh[j]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            is_bold = (j == 0)
            color = (0x1E, 0x3A, 0x8A) if is_bold else (0x33, 0x41, 0x55)
            format_cell_text(cell, val, bold=is_bold, color_rgb=color, size=7.0)
            
    doc.add_page_break()
    
    # ================= PAGE 3 =================
    add_styled_heading(doc, "7. System Working & Dispatch Logic", level=1)
    
    add_styled_heading(doc, "7.1 EMS Power Flow State Machine", level=2)
    p_flow = doc.add_paragraph()
    p_flow.paragraph_format.space_after = Pt(6)
    p_flow.paragraph_format.line_spacing = 1.15
    run_flow = p_flow.add_run(
        "At each second, net surplus is evaluated: P_surplus = P_solar + P_wind - P_load. The dispatcher routes flows based on priority rules:\n"
        "1. Renewables Priority: PV and Wind are routed directly to satisfy Load demand first.\n"
        "2. Renewable Surplus (P_surplus >= 0): Surplus power is directed to charge the battery bank. If Battery SOC >= 95%, the surplus is exported to the Grid.\n"
        "3. Renewable Deficit (P_surplus < 0): If Battery SOC > 20%, the battery discharges to cover the load. Grid import is used as a final fallback.\n"
        "4. Islanding Mode (Grid Outage): Utility grid status drops to 0. Breakers open. Tiered load shedding initiates based on SOC:\n"
        "   - SOC <= 25% (Level 1): Shed 30% load (isolate non-critical facility blocks).\n"
        "   - SOC <= 20% (Level 2): Shed 60% load (isolate heavy HVAC machinery).\n"
        "   - SOC <= 15% (Level 3): Shed 85% load (vital infrastructure hospital/safety lines only)."
    )
    run_flow.font.name = 'Calibri'
    run_flow.font.size = Pt(8.5)
    run_flow.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # Add Visual Flowchart Table directly into the document
    add_styled_heading(doc, "Microgrid Controller 1Hz Execution Flowchart", level=2)
    add_flowchart_table(doc)
    
    add_styled_heading(doc, "7.2 Q-Learning Tariff Optimizer", level=2)
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_after = Pt(6)
    p_q.paragraph_format.line_spacing = 1.15
    run_q = p_q.add_run(
        "A reinforcement learning agent schedules charge/discharge offsets during grid-connected modes to minimize costs:\n"
        "- Discretized States: S = (hour_bin, soc_bin, surplus_bin) where hour is grouped in 4-hour bins, SOC in 20% bins, and surplus into [Deficit, Balanced, Surplus].\n"
        "- Discrete Actions: A = [-20 kW (Discharge), 0 kW (Standby), +20 kW (Charge)].\n"
        "- Time-of-Use Tariff rates: ₹7.50 / kWh peak (14:00 - 20:00), and ₹4.50 / kWh off-peak.\n"
        "- Reward Function: Reward = -(Electricity Cost) - beta * |I_battery|, where beta represents a cell degradation cycling penalty.\n"
        "- Bellman Q-Value updates: Q(s,a) = Q(s,a) + alpha * [Reward + gamma * max_a(Q(s', a')) - Q(s,a)]."
    )
    run_q.font.name = 'Calibri'
    run_q.font.size = Pt(8.5)
    run_q.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    add_styled_heading(doc, "7.3 Isolation Forest Anomaly Detection", level=2)
    p_anom = doc.add_paragraph()
    p_anom.paragraph_format.space_after = Pt(6)
    run_anom = p_anom.add_run(
        "Identifies sensor deviations or dropouts by evaluating the feature vector: X = [Grid_Voltage, Load_Current, Battery_Temperature, Inverter_Efficiency] "
        "against an isolation model trained with a Contamination factor of 3%. Outlier scores closer to -1.0 trigger alarm 'AI-ANOM-001' with troubleshooting steps."
    )
    run_anom.font.name = 'Calibri'
    run_anom.font.size = Pt(8.5)
    run_anom.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    add_styled_heading(doc, "7.4 Predictive Maintenance RUL Formulas", level=2)
    p_rul = doc.add_paragraph()
    p_rul.paragraph_format.space_after = Pt(6)
    p_rul.paragraph_format.line_spacing = 1.15
    run_rul = p_rul.add_run(
        "- PCS Inverter Health: H_inv = 98.2 - 1.5 * max(0, (T_inverter - 50.0) / 30.0). Inverter RUL = 12000 hours * (H_inv / 100).\n"
        "- Battery SOH thermal aging: Degradation is exponential to temperature: Kappa = exp(0.04 * (T_battery - 25.0)). BESS RUL = (6000 * SOH / 100) / Kappa."
    )
    run_rul.font.name = 'Calibri'
    run_rul.font.size = Pt(8.5)
    run_rul.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # 8. Backend & Frontend Code Design
    add_styled_heading(doc, "8. Backend & Frontend Code Design", level=1)
    
    code_headers = ["File Name", "Component / Functions", "Operational Description"]
    code_rows = [
        ["backend/main.py", "FastAPI App Server\nrun_microgrid_loop()\nws_endpoint()", "Runs non-blocking 1Hz daemon. Loosely reads CSV datasets in fallback, processes active WebSockets telemetry streams, and services REST overrides."],
        ["backend/database.py", "SQLite Database Manager\nlog_telemetry()\ntrigger_alarm()", "Logs 1Hz historian metrics to relational tables. Prevents alarm database duplication spam using message prefix-matching rules."],
        ["backend/ems.py", "EMSEngine Core\ncompute_dispatch()\ndetect_alarms()", "Executes rule state machine, derates charge currents based on thermal stress, and matches alarms with field repair guides."],
        ["backend/ai_models.py", "AI & Forecasting Suite\nLSTMLoadForecaster\nAnomalyDetector", "MLP neural networks forecast loads. Isolation Forest classifies outliers. Computes component degradation & RUL calculations."],
        ["backend/protocols.py", "Register Brokers\nModbusTCPSimulator\nCANBusSimulator", "Maps metrics to Holding registers (40001+). Packs voltage, current, and temp bytes into 8-byte frames at CAN ID 0x18F009A1."],
        ["backend/simulator.py", "Physics Twin Simulation\nupdate_battery_physics()", "Simulates dynamic ambient walks, panel irradiance heat, battery cell internal resistance voltage drop, and fan active cooling flows."],
        ["frontend/src/App.jsx", "React HMI Panel View", "Responsive dashboard containing dynamic SVG synoptic flows (animated speeds), Data Sources manager, Historian logger, and alarm acknowledge buttons."],
        ["index.js / html", "Vanilla HMI client", "Legacy local dashboard mapping SVG particles and canvas trend graphs."]
    ]
    
    code_table = doc.add_table(rows=len(code_rows) + 1, cols=3)
    code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(code_table)
    
    col_widths_code = [Inches(1.5), Inches(2.2), Inches(3.8)]
    
    hdr_row = code_table.rows[0]
    for j, text in enumerate(code_headers):
        cell = hdr_row.cells[j]
        cell.width = col_widths_code[j]
        set_cell_shading(cell, "1E3A8A")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        format_cell_text(cell, text, bold=True, color_rgb=(0xFF, 0xFF, 0xFF), size=8.0)
        
    for i, row_data in enumerate(code_rows):
        row = code_table.rows[i + 1]
        bg_color = "FFFFFF" if i % 2 == 0 else "F8FAFC"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths_code[j]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            is_bold = (j == 0)
            color = (0x1E, 0x3A, 0x8A) if is_bold else (0x33, 0x41, 0x55)
            format_cell_text(cell, val, bold=is_bold, color_rgb=color, size=7.0)
            
    doc.add_page_break()
    
    # ================= PAGE 4 =================
    add_styled_heading(doc, "9. Database Schema & Mappings", level=1)
    
    schema_headers = ["Table Name", "Columns Schema", "Mapping Description"]
    schema_rows = [
        ["ems_settings", "key (TEXT PRIMARY KEY), value (TEXT)", "Stores configuration setpoints (e.g. battery_min_soc: 20.0, tariff_peak_rate: 7.50)."],
        ["alarms_log", "id (INTEGER PK AUTOINCREMENT), timestamp (REAL), severity (TEXT), source (TEXT), message (TEXT), status (TEXT)", "Chronological alarms journal. Status values: ACTIVE, ACKNOWLEDGED, CLEARED."],
        ["scada_historian", "timestamp (REAL PRIMARY KEY), solar_power (REAL), solar_voltage (REAL), wind_power (REAL), battery_soc (REAL), battery_soh (REAL), battery_voltage (REAL), battery_current (REAL), battery_temperature (REAL), grid_status (INTEGER), grid_voltage (REAL), grid_frequency (REAL), grid_power (REAL), load_demand (REAL), load_current (REAL), load_voltage (REAL), inverter_status (TEXT), inverter_efficiency (REAL), inverter_output_power (REAL), ems_action (TEXT), electricity_cost (REAL)", "Stores 1Hz historical telemetry points of the microgrid system. Used for off-line diagnostics and analytical performance audits."],
        ["solar_realtime", "timestamp (REAL PRIMARY KEY), irradiance, ambient_temp, panel_temp, dc_voltage, dc_current, ac_power, inverter_status...", "Stores high-frequency sensor readings specifically connected to the Solar generation stack."],
        ["wind_realtime", "timestamp (REAL PRIMARY KEY), wind_speed, wind_direction, air_density, blade_angle, turbine_rpm, generator_voltage...", "Stores high-frequency sensor readings connected to the Wind turbine generation stack."],
        ["battery_realtime", "timestamp (REAL PRIMARY KEY), soc, soh, voltage, current, temperature, cell_voltage, cell_temperature, charge_rate...", "Stores high-frequency BMS readings connected to the BESS stack."],
        ["grid_realtime", "timestamp (REAL PRIMARY KEY), voltage, current, frequency, power_factor, import_power, export_power", "Stores grid voltage, frequencies and importing/exporting active power stats."],
        ["load_realtime", "timestamp (REAL PRIMARY KEY), load_voltage, load_current, active_power, reactive_power, apparent_power...", "Stores active/reactive power readings and feeder energy consumption totals."]
    ]
    
    schema_table = doc.add_table(rows=len(schema_rows) + 1, cols=3)
    schema_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(schema_table)
    
    col_widths_schema = [Inches(1.5), Inches(3.2), Inches(2.8)]
    
    hdr_row = schema_table.rows[0]
    for j, text in enumerate(schema_headers):
        cell = hdr_row.cells[j]
        cell.width = col_widths_schema[j]
        set_cell_shading(cell, "1E3A8A")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        format_cell_text(cell, text, bold=True, color_rgb=(0xFF, 0xFF, 0xFF), size=8.0)
        
    for i, row_data in enumerate(schema_rows):
        row = schema_table.rows[i + 1]
        bg_color = "FFFFFF" if i % 2 == 0 else "F8FAFC"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths_schema[j]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            is_bold = (j == 0)
            color = (0x1E, 0x3A, 0x8A) if is_bold else (0x33, 0x41, 0x55)
            format_cell_text(cell, val, bold=is_bold, color_rgb=color, size=7.0)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # 10. Implementation Timeline
    add_styled_heading(doc, "10. Implementation Timeline", level=1)
    
    timeline_headers = ["Phase", "Milestone", "Deliverables", "Status"]
    timeline_rows = [
        ["P1", "Infrastructure setup", "FastAPI web server layout, SQLite database schemas mapping.", "Completed"],
        ["P2", "Responsive Web HMI", "Dynamic SVG synoptic flow animations, live canvas gauges.", "Completed"],
        ["P3", "Settings Management", "Relational configuration setpoints, peak/off-peak pricing overrides API.", "Completed"],
        ["P4", "Modbus TCP Protocol", "Holding registers mapping (40001 - 40010) and scaling equations.", "Completed"],
        ["P5", "CAN Bus BMS Gateway", "8-byte hexadecimal telemetry package framing, checksums.", "Completed"],
        ["P6", "AI Q-Learning Agent", "State space discretizations, Bellman policy updates, tariff shifts reward models.", "Completed"],
        ["P7", "OPC UA & IEC Nodes", "Node set directories, Circuit Breakers logical node statuses.", "Completed"],
        ["P8", "Fallback Replay Loop", "Caching dataset lists at boot, 1Hz sequential row looping.", "Completed"],
        ["P9", "Report Compilation", "Unit test validations, DOCX compilation, technical report.", "Completed"]
    ]
    
    timeline_table = doc.add_table(rows=len(timeline_rows) + 1, cols=4)
    timeline_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(timeline_table)
    
    col_widths_time = [Inches(0.6), Inches(1.8), Inches(3.9), Inches(1.2)]
    
    hdr_row = timeline_table.rows[0]
    for j, text in enumerate(timeline_headers):
        cell = hdr_row.cells[j]
        cell.width = col_widths_time[j]
        set_cell_shading(cell, "1E3A8A")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        format_cell_text(cell, text, bold=True, color_rgb=(0xFF, 0xFF, 0xFF), size=8.0)
        
    for i, row_data in enumerate(timeline_rows):
        row = timeline_table.rows[i + 1]
        bg_color = "FFFFFF" if i % 2 == 0 else "F8FAFC"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths_time[j]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            is_bold = (j == 0)
            if j == 3:
                color_rgb = (0x06, 0x5F, 0x46)
            else:
                color_rgb = (0x1E, 0x3A, 0x8A) if is_bold else (0x33, 0x41, 0x55)
            format_cell_text(cell, val, bold=is_bold or (j==3), color_rgb=color_rgb, size=7.0)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Conclusion Styled Box
    conc_table = doc.add_table(rows=1, cols=1)
    conc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    conc_cell = conc_table.rows[0].cells[0]
    conc_cell.width = Inches(7.5)
    set_cell_shading(conc_cell, "EFF6FF")
    set_cell_margins(conc_cell, top=140, bottom=140, left=180, right=180)
    
    tcPr = conc_cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'BFDBFE')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    p_conc = conc_cell.paragraphs[0]
    p_conc.paragraph_format.space_before = Pt(0)
    p_conc.paragraph_format.space_after = Pt(0)
    p_conc.paragraph_format.line_spacing = 1.15
    
    run_c_hdr = p_conc.add_run("Conclusion\n")
    run_c_hdr.font.name = 'Calibri'
    run_c_hdr.font.bold = True
    run_c_hdr.font.size = Pt(8.5)
    run_c_hdr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    run_c_body = p_conc.add_run(
        "APEX is a modular, high-fidelity microgrid SCADA and Energy Management System platform. By integrating "
        "industrial communication protocols (Modbus, CAN Bus, OPC UA, IEC 61850) with automated dispatch logic and "
        "reinforcement learning cost optimization, it achieves significant operational utility cost reductions while "
        "protecting BESS storage assets. The dual HMI client web panels and resilient cached fallback loops provide "
        "robust, reliable supervisory control and time-series telemetry data acquisition suitable for modern microgrid utilities."
    )
    run_c_body.font.name = 'Calibri'
    run_c_body.font.size = Pt(8.5)
    run_c_body.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    doc.save("APEX_Microgrid_SCADA_EMS_Detailed_System_Report_Final.docx")
    print("Word Document generated successfully: APEX_Microgrid_SCADA_EMS_Detailed_System_Report_Final.docx")

if __name__ == "__main__":
    build_docx_report()
